from rest_framework.decorators import api_view,permission_classes
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from rest_framework import status
import json
from .scorestudent import getTestSeriesQuestions, getMCQs, parse_json_from_gemini
from langchain_groq import ChatGroq
import os
from .models import PYQSubject, PYQquestions, Subject
from .serializers import PYQSubjectSerialiser
from django.db.models import Count
from django.core.cache import cache
from django.views.decorators.cache import cache_page
from reportlab.pdfgen import canvas 
from reportlab.pdfbase.ttfonts import TTFont 
from reportlab.pdfbase import pdfmetrics 
from reportlab.lib import colors 
from docx import Document as docMaker
from docx2pdf import convert

groq_api_key=os.environ['GROQ_API_KEY']


@api_view(['GET'])
def get_test_subjects(request): 
    subjects = PYQSubject.objects.all()
    serializer = PYQSubjectSerialiser(subjects, many=True)
    return Response(serializer.data)

@api_view(['GET'])
def pyq_by_year_subject_view(request, subject_id):
    if request.method == 'GET':
        grouped_questions = PYQquestions.objects.filter(subject_id=subject_id).values('year').annotate(total_questions=Count('id')).order_by('year')

        # Prepare response data
        data = []
        for entry in grouped_questions:
            year = entry['year']
            total_questions = entry['total_questions']
            questions = PYQquestions.objects.filter(subject_id=subject_id, year=year).values('question', 'marks')
            data.append({
                'year': year,
                'total_questions': total_questions,
                'questions': list(questions)
            })

        return Response(data)
    
@api_view(['GET'])
@cache_page(60 * 60)
def generateTest(request, subject_id):
    cached_score = cache.get(subject_id)
    if cached_score:
        return Response(cached_score, status=status.HTTP_200_OK)
    
    if request.method == 'GET':
        llm=ChatGroq(groq_api_key=groq_api_key,
             model_name="llama3-8b-8192")
        all_questions = PYQquestions.objects.filter(subject_id=subject_id)
        ten_marks_q = []
        five_marks_q = []
        test = []
        for question in all_questions:
            if question.marks == 5:
                five_marks_q.append(question.question)
            else:
                ten_marks_q.append(question.question,)
        print("true")
        five_marks = getTestSeriesQuestions(llm,five_marks_q,250)
        ctr =0
        while not five_marks and ctr < 4:
            print(ctr)
            five_marks = getTestSeriesQuestions(llm,five_marks_q,250)
            ctr +=1
        test.append(
            {
                "questions" : five_marks.get('solutions',[]),
                "marks" : 5
            })
        print("appended 5 marks")

        ten_marks = getTestSeriesQuestions(llm,ten_marks_q,500)
        ctr = 0
        while not ten_marks and ctr < 4:
            print(ctr)
            ten_marks = getTestSeriesQuestions(llm,ten_marks_q,500)
            ctr +=1
        # print(ten_marks)
        test.append(
            {
                "questions" : ten_marks.get('solutions',[]),
                "marks" : 10
            })
        print("appended 10 marks")
        subject = PYQSubject.objects.get(pk=subject_id)
        x = getMCQs(llm, subject.name)
        test.append({"questions": x.get('solutions',[]),
                    "marks" : 1})

        cache.set(subject_id, test, timeout=None)
        return Response(test)


@api_view(['POST'])
def pdfTestSeriesQuestions(request,subject_id):
    data=request.data
    

    document = docMaker()
    document.add_heading('Test Series', 0)
    for section in data:
        document.add_heading(f"{section['marks']} mark questions", level=1)
        ctr=1
        for question_data in section['questions']:
            document.add_paragraph(f"{ctr}). {question_data['question']}")
            ctr+=1
            if section['marks'] == 1:  
                for idx, option in enumerate(question_data['options']):
                    document.add_paragraph(f"{chr(97 + idx)}. {option}")

    document.save(f'./study/mediafiles/{subject_id}.docx')

    convert(f'./study/mediafiles/{subject_id}.docx', f'./study/mediafiles/{subject_id}.pdf')
    os.remove(f'./study/mediafiles/{subject_id}.docx')

    return Response(
        {
            "file": f'/media/{subject_id}.pdf'
        }
    )

